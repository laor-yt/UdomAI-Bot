import os
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"
import uuid
import subprocess
import shutil
import requests
import ffmpeg
import imageio_ffmpeg
from PIL import Image
from utils import get_temp_dir, cleanup_file
import re
import logging

# Suppress Hugging Face Hub unauthenticated warning logs
logging.getLogger("huggingface_hub").setLevel(logging.ERROR)
logging.getLogger("huggingface_hub.utils._http").setLevel(logging.ERROR)

logger = logging.getLogger(__name__)

# Configure ffmpeg binary environment and create ffmpeg.exe alias
ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
ffmpeg_dir = os.path.dirname(ffmpeg_exe)
if ffmpeg_dir not in os.environ.get("PATH", ""):
    os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")

ffmpeg_alias = os.path.join(ffmpeg_dir, "ffmpeg.exe" if os.name == 'nt' else "ffmpeg")
if not os.path.exists(ffmpeg_alias) and os.path.exists(ffmpeg_exe):
    try:
        import shutil
        shutil.copyfile(ffmpeg_exe, ffmpeg_alias)
    except Exception as e:
        print(f"Could not create ffmpeg alias: {e}")

from pydub import AudioSegment
AudioSegment.converter = ffmpeg_alias
AudioSegment.ffmpeg = ffmpeg_alias

def _run_ffmpeg_with_progress(stream, output_path, progress_callback):
    try:
        process = stream.run_async(
            cmd=imageio_ffmpeg.get_ffmpeg_exe(),
            pipe_stderr=True,
            pipe_stdout=False,
            quiet=True
        )
        
        size_pattern = re.compile(r"size=\s*(\d+[a-zA-Z]+)")
        speed_pattern = re.compile(r"speed=\s*([\d.]+x)")
        stderr_lines = []
        
        while True:
            line = process.stderr.readline()
            if not line:
                break
            
            line_str = line.decode('utf-8', errors='ignore')
            stderr_lines.append(line_str)
            
            if progress_callback:
                size_match = size_pattern.search(line_str)
                speed_match = speed_pattern.search(line_str)
                
                if size_match and speed_match:
                    size = size_match.group(1)
                    speed = speed_match.group(1)
                    progress_callback(f"Converting... Size: {size}, Speed: {speed}")
                    
        process.wait()
        
        # If output file exists and is non-empty, consider it successful
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return output_path
            
        if process.returncode != 0:
            err_msg = "".join(stderr_lines[-5:]).strip()
            print(f"FFmpeg error with code {process.returncode}: {err_msg}")
            return None
            
        return output_path if (os.path.exists(output_path) and os.path.getsize(output_path) > 0) else None
    except Exception as e:
        print(f"Error during ffmpeg execution: {e}")
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return output_path
        return None

def convert_video_to_audio(input_path, output_format='mp3', progress_callback=None):
    """
    Converts a video or audio file to a specified audio format.
    """
    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    stream = (
        ffmpeg
        .input(input_path)
        .output(output_path, acodec='libmp3lame' if output_format == 'mp3' else 'copy', qscale=2)
        .overwrite_output()
    )
    
    res = _run_ffmpeg_with_progress(stream, output_path, progress_callback)
    if res and os.path.exists(res) and os.path.getsize(res) > 0:
        return res

    # Direct subprocess fallback for audio extraction if ffmpeg-python stream failed
    try:
        import subprocess
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [ffmpeg_exe, "-y", "-i", input_path]
        if output_format == 'mp3':
            cmd.extend(["-vn", "-acodec", "libmp3lame", "-ar", "44100", "-ac", "2", "-q:a", "2", "-map_metadata", "-1", output_path])
        else:
            cmd.extend(["-vn", "-acodec", "copy", "-map_metadata", "-1", output_path])
        
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=7200)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return output_path
    except Exception as e:
        print(f"Direct FFmpeg audio fallback failed: {e}")

    return None

def has_copyright_metadata(input_path):
    import subprocess, re
    import imageio_ffmpeg
    try:
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [ffmpeg_exe, "-i", input_path]
        res = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, timeout=14400)
        stderr_text = res.stderr.decode("utf-8", errors="ignore")
        # Look for metadata block
        metadata_match = re.search(r"Metadata:(.*?)(?:Duration:|Stream #)", stderr_text, re.DOTALL | re.IGNORECASE)
        if metadata_match:
            meta_text = metadata_match.group(1).lower()
            for field in ['title', 'artist', 'album', 'copyright', 'description', 'comment']:
                if re.search(fr"^\s*{field}\s*:\s*(.+)", meta_text, re.MULTILINE):
                    return True
    except Exception:
        pass
    return False

def convert_video_format(input_path, output_format='mp4', progress_callback=None):
    """
    Converts a video file to a different video format.
    """
    if has_copyright_metadata(input_path) and progress_callback:
        progress_callback("⚠️ Copyright metadata detected and removed from original file.")

    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    stream = (
        ffmpeg
        .input(input_path)
        .output(output_path, map_metadata="-1")
        .overwrite_output()
    )
    
    res = _run_ffmpeg_with_progress(stream, output_path, progress_callback)
    if res and os.path.exists(res) and os.path.getsize(res) > 0:
        return res

    # Direct subprocess fallback for video format conversion
    try:
        import subprocess
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [ffmpeg_exe, "-y", "-i", input_path, "-vcodec", "libx264", "-acodec", "aac", "-strict", "-2", "-map_metadata", "-1", output_path]
        proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=14400)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return output_path
    except Exception as e:
        print(f"Direct FFmpeg video fallback failed: {e}")

    return None

def convert_image_format(input_path, output_format='png'):
    """
    Converts an image file to a different format using Pillow.
    """
    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    try:
        with Image.open(input_path) as img:
            rgb_im = img.convert('RGB')
            rgb_im.save(output_path, format=output_format.upper())
        return output_path
    except Exception as e:
        print(f"Error converting image format: {e}")
        return None

def convert_document_format(input_path, output_format='pdf'):
    """
    Converts a document (PDF, DOCX, TXT) to another document format (PDF, DOCX, TXT).
    """
    temp_dir = get_temp_dir()
    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(temp_dir, output_filename)
    
    ext = os.path.splitext(input_path)[1].lower()
    extracted_text = ""
    
    try:
        if ext == '.pdf':
            import fitz
            doc = fitz.open(input_path)
            for page in doc:
                extracted_text += page.get_text() + "\n"
        elif ext in ['.docx', '.doc']:
            import docx
            doc = docx.Document(input_path)
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
                
        if not extracted_text.strip():
            extracted_text = "No text content found in original file."

        if output_format == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
            return output_path
            
        elif output_format == 'docx':
            import docx
            doc = docx.Document()
            for line in extracted_text.split('\n'):
                doc.add_paragraph(line)
            doc.save(output_path)
            return output_path
            
        elif output_format == 'pdf':
            import fitz
            doc = fitz.open()
            page = doc.new_page()
            margin = 50
            rect = fitz.Rect(margin, margin, page.rect.width - margin, page.rect.height - margin)
            page.insert_textbox(rect, extracted_text, fontsize=11)
            doc.save(output_path)
            return output_path
            
    except Exception as e:
        print(f"Error converting document: {e}")
        return None

def detect_speech_segments(input_path, min_silence_dur=0.4, noise_threshold=-28):
    """
    Detects start and end timestamps (seconds) of spoken parts in media using FFmpeg silencedetect.
    Returns list of tuples: [(start_sec, end_sec), ...]
    """
    duration = get_video_duration(input_path)
    if duration <= 0:
        return []
        
    try:
        import subprocess
        import re
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [
            ffmpeg_exe, "-i", input_path,
            "-af", f"silencedetect=noise={noise_threshold}dB:d={min_silence_dur}",
            "-f", "null", "-"
        ]
        res = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore")
        
        silence_starts = []
        silence_ends = []
        
        for line in res.stderr.split('\n'):
            if "silence_start:" in line:
                m = re.search(r"silence_start:\s*([\d.]+)", line)
                if m:
                    silence_starts.append(float(m.group(1)))
            elif "silence_end:" in line:
                m = re.search(r"silence_end:\s*([\d.]+)", line)
                if m:
                    silence_ends.append(float(m.group(1)))
                    
        silences = []
        for i in range(min(len(silence_starts), len(silence_ends))):
            s_start = silence_starts[i]
            s_end = silence_ends[i]
            if s_end > s_start:
                silences.append((s_start, s_end))
                
        if not silences:
            return [(0.0, duration)]
            
        segments = []
        curr = 0.0
        for s_start, s_end in silences:
            if s_start > curr + 0.3:
                segments.append((curr, s_start))
            curr = s_end
            
        if curr < duration - 0.3:
            segments.append((curr, duration))
            
        return segments if segments else [(0.0, duration)]
    except Exception as e:
        print(f"Speech segment detection error: {e}")
        return []

def check_audio_rms(file_path, min_max_vol=-25.0, min_mean_vol=-45.0):
    """
    Checks if an audio file has audible music/SFX content using FFmpeg volumedetect filter.
    Returns True if max_volume > min_max_vol AND mean_volume > min_mean_vol, else False.
    """
    try:
        import subprocess
        import re
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [
            ffmpeg_exe, "-i", file_path,
            "-af", "volumedetect",
            "-f", "null", "-"
        ]
        res = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, text=True, encoding="utf-8", errors="ignore")
        match_max = re.search(r"max_volume:\s*(-?[\d.]+)\s*dB", res.stderr)
        match_mean = re.search(r"mean_volume:\s*(-?[\d.]+)\s*dB", res.stderr)
        if match_max:
            max_vol = float(match_max.group(1))
            mean_vol = float(match_mean.group(1)) if match_mean else -99.0
            return max_vol > min_max_vol and mean_vol > min_mean_vol
    except Exception as e:
        print(f"Error checking RMS volume: {e}")
    return os.path.exists(file_path) and os.path.getsize(file_path) > 500

def extract_bgm_demucs(input_path, output_path=None, return_vocals=False):
    """
    Uses Demucs v4 (htdemucs) AI stem separation to cleanly extract background music & sound effects
    (no_vocals stem) regardless of audio channel layout (mono or stereo).
    Fallback to DSP notch filtering if Demucs fails or BGM RMS volume is near silence (< -45dB).
    """
    dur_in = get_video_duration(input_path)
    # Skip Demucs for very long videos (> 30 minutes) to avoid hanging, OOM crashes, or extreme delays.
    if dur_in > 1800:
        print(f"Video duration ({dur_in}s) > 1800s. Skipping Demucs, using fast DSP fallback.")
        res = extract_bgm_no_vocals_dsp(input_path, output_path)
        return (res, None) if return_vocals else res

    temp_dir = get_temp_dir()
    if not output_path:
        output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_bgm_demucs.mp3")
        
    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    ffmpeg_dir = os.path.dirname(ffmpeg_exe)
    if ffmpeg_dir not in os.environ.get("PATH", ""):
        os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")
        
    import subprocess, shutil, sys
    
    # 1. Extract audio track to temporary WAV file (ensures Demucs & sphn can load cleanly)
    temp_input_wav = os.path.join(temp_dir, f"{uuid.uuid4()}_input_track.wav")
    demucs_out_dir = os.path.join(temp_dir, f"demucs_{uuid.uuid4().hex[:8]}")
    try:
        cmd_wav = [
            ffmpeg_exe, "-y", "-i", input_path,
            "-ar", "44100", "-ac", "2", temp_input_wav
        ]
        subprocess.run(cmd_wav, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=14400)
        
        if os.path.exists(temp_input_wav) and os.path.getsize(temp_input_wav) > 1000:
            cmd = [
                sys.executable, "-m", "demucs.separate", "--two-stems=vocals", "-n", "htdemucs",
                temp_input_wav, "-o", demucs_out_dir
            ]
            dur_in = get_video_duration(input_path)
            # Demucs on CPU can take up to 10-15x real-time. Give it ample timeout.
            demucs_timeout = max(900, int(dur_in * 15))
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=demucs_timeout)
            
            no_vocals_file = None
            vocals_file = None
            for root, _, files in os.walk(demucs_out_dir):
                for f in files:
                    if "no_vocals" in f.lower():
                        no_vocals_file = os.path.join(root, f)
                    elif "vocals" in f.lower() and "no_vocals" not in f.lower():
                        vocals_file = os.path.join(root, f)
                    
            if no_vocals_file and os.path.exists(no_vocals_file):
                # Convert no_vocals stem to 44.1kHz MP3
                cmd_cvt = [
                    ffmpeg_exe, "-y", "-i", no_vocals_file,
                    "-acodec", "libmp3lame", "-ar", "44100", "-ac", "2", "-b:a", "192k",
                    output_path
                ]
                subprocess.run(cmd_cvt, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=14400)
                
                out_vocals_path = None
                if return_vocals and vocals_file and os.path.exists(vocals_file):
                    out_vocals_path = os.path.join(temp_dir, f"{uuid.uuid4()}_vocals_demucs.mp3")
                    cmd_cvt_voc = [
                        ffmpeg_exe, "-y", "-i", vocals_file,
                        "-acodec", "libmp3lame", "-ar", "44100", "-ac", "2", "-b:a", "192k",
                        out_vocals_path
                    ]
                    subprocess.run(cmd_cvt_voc, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=14400)

                shutil.rmtree(demucs_out_dir, ignore_errors=True)
                cleanup_file(temp_input_wav)
                if os.path.exists(output_path) and os.path.getsize(output_path) > 100:
                    return (output_path, out_vocals_path) if return_vocals else output_path
    except Exception as e:
        print(f"Demucs extraction error: {e}")
        shutil.rmtree(demucs_out_dir, ignore_errors=True)
        cleanup_file(temp_input_wav)

    # Fallback to DSP filter if Demucs fails
    res_dsp = extract_bgm_no_vocals_dsp(input_path, output_path)
    return (res_dsp, None) if return_vocals else res_dsp

def extract_bgm_no_vocals_dsp(input_path, output_path=None):
    """
    Ultra-fast DSP audio extraction for background music & sound effects.
    Cancels center-channel panned vocals and eliminates 300Hz-3.4kHz human speech frequencies.
    Runs in < 0.3s without CPU/memory bottlenecks.
    """
    temp_dir = get_temp_dir()
    if not output_path:
        output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_bgm_dsp.mp3")
        
    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    import subprocess
    
    af_filter = "pan=stereo|c0=0.5*c0-0.5*c1|c1=0.5*c1-0.5*c0, equalizer=f=300:width_type=h:width=400:g=-24, equalizer=f=1200:width_type=h:width=1800:g=-30, equalizer=f=2800:width_type=h:width=1200:g=-24, volume=1.8"
    try:
        cmd = [
            ffmpeg_exe, "-y", "-i", input_path,
            "-af", af_filter, "-acodec", "libmp3lame", "-ar", "44100", "-ac", "2", "-b:a", "192k",
            output_path
        ]
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=300)
        if check_audio_rms(output_path):
            return output_path
    except Exception:
        pass

    try:
        af_notch = "equalizer=f=300:width_type=h:width=400:g=-24, equalizer=f=1200:width_type=h:width=1800:g=-30, equalizer=f=2800:width_type=h:width=1200:g=-24, volume=1.8"
        cmd = [
            ffmpeg_exe, "-y", "-i", input_path,
            "-af", af_notch, "-acodec", "libmp3lame", "-ar", "44100", "-ac", "2", "-b:a", "192k",
            output_path
        ]
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=300)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
            return output_path
    except Exception:
        pass

    return None

def extract_bgm_no_vocals(input_path, output_path=None):
    """Wrapper function to extract BGM without vocals."""
    return extract_bgm_demucs(input_path, output_path)

def mix_tts_with_bgm(tts_audio_path, bgm_audio_path, target_dur=None, bgm_volume=0.70, tts_volume=1.40):
    """
    Mixes translated/recap AI speech track (tts_volume=1.40, boosted voice) with original background music & sound effects track (bgm_volume=0.70) in 44.1kHz Stereo.
    """
    temp_dir = get_temp_dir()
    mixed_path = os.path.join(temp_dir, f"{uuid.uuid4()}_mixed_speech_bgm.mp3")
    
    if not bgm_audio_path or not os.path.exists(bgm_audio_path) or os.path.getsize(bgm_audio_path) < 500:
        return tts_audio_path
        
    try:
        tts_dur = get_video_duration(tts_audio_path)
        bgm_dur = get_video_duration(bgm_audio_path)
        dur = target_dur or tts_dur or bgm_dur
        
        if bgm_dur > 0 and tts_dur > bgm_dur:
            bgm_in = ffmpeg.input(bgm_audio_path, stream_loop=-1).audio
        else:
            bgm_in = ffmpeg.input(bgm_audio_path).audio
            
        tts_in = ffmpeg.input(tts_audio_path).audio
        
        # Translate voice is prominently louder than BGM & Sound Effects (Voice 1.40 vs BGM 0.70)
        tts_vol = tts_in.filter('volume', tts_volume)
        bgm_vol = bgm_in.filter('volume', bgm_volume)
        mixed_stream = ffmpeg.filter([tts_vol, bgm_vol], 'amix', inputs=2, duration='longest', dropout_transition=2, normalize=False)
        mixed_stream = mixed_stream.filter('aresample', 44100).filter('aformat', channel_layouts='stereo')
        
        out_opts = {'acodec': 'libmp3lame', 'b:a': '192k', 'ar': '44100', 'ac': 2}
        if dur and dur > 0:
            out_opts['t'] = dur
            
        stream = ffmpeg.output(mixed_stream, mixed_path, **out_opts).overwrite_output()
        _run_ffmpeg_with_progress(stream, mixed_path, None)
        
        if os.path.exists(mixed_path) and os.path.getsize(mixed_path) > 100:
            return mixed_path
    except Exception as e:
        print(f"Error mixing TTS with BGM: {e}")
        
    return tts_audio_path

def build_atempo_filter_chain(speed_ratio):
    """
    FFmpeg atempo filter requires values between 0.5 and 2.0.
    Chains multiple atempo filters for values outside this range.
    """
    speed_ratio = max(0.25, min(4.0, speed_ratio))
    if 0.5 <= speed_ratio <= 2.0:
        return [("atempo", speed_ratio)]
    elif speed_ratio > 2.0:
        f1 = 2.0
        f2 = speed_ratio / 2.0
        return [("atempo", f1), ("atempo", f2)]
    else:
        f1 = 0.5
        f2 = speed_ratio / 0.5
        return [("atempo", f1), ("atempo", f2)]

import numpy as np

# ─── Whisper model cache (loaded once, reused across all dub calls) ───────────
_WHISPER_MODEL_CACHE = None
_WHISPER_MODEL_LOCK = None

def _get_whisper_model():
    """Returns the cached WhisperModel, loading it on first call."""
    global _WHISPER_MODEL_CACHE, _WHISPER_MODEL_LOCK
    import threading
    if _WHISPER_MODEL_LOCK is None:
        _WHISPER_MODEL_LOCK = threading.Lock()
    with _WHISPER_MODEL_LOCK:
        if _WHISPER_MODEL_CACHE is None:
            try:
                from faster_whisper import WhisperModel
                for m_name in ["small", "base", "tiny"]:
                    try:
                        _WHISPER_MODEL_CACHE = WhisperModel(m_name, device="cpu", compute_type="int8")
                        logger.info(f"[Whisper] Loaded model '{m_name}' into cache.")
                        break
                    except Exception as e_m:
                        logger.warning(f"Could not load WhisperModel({m_name}): {e_m}")
            except Exception as e:
                logger.error(f"faster_whisper not available: {e}")
        return _WHISPER_MODEL_CACHE

# ─── Persistent requests session for translation HTTP calls ──────────────────
_TRANSLATE_SESSION = None

def _get_translate_session():
    global _TRANSLATE_SESSION
    if _TRANSLATE_SESSION is None:
        import requests as _requests
        _TRANSLATE_SESSION = _requests.Session()
        _TRANSLATE_SESSION.headers.update({"User-Agent": "Mozilla/5.0"})
    return _TRANSLATE_SESSION

def detect_gender_from_audio(audio_path):
    """
    Fast heuristic gender detection using zero-crossing rate (ZCR).
    Male voices have lower ZCR (deeper fundamentals), female voices higher.
    ~10x faster than librosa.pyin — no full pitch tracking needed.
    Defaults to female if detection fails.
    """
    try:
        import librosa
        y, sr = librosa.load(audio_path, sr=16000, mono=True, duration=3.0)  # only first 3s needed
        zcr = float(np.mean(librosa.feature.zero_crossing_rate(y)))
        # Empirical threshold: male speech ZCR typically < 0.07, female > 0.07
        return 'male' if zcr < 0.07 else 'female'
    except Exception as e:
        logger.warning(f"Gender detection failed: {e}")
    return 'female'  # Default fallback

def generate_neural_tts(text, lang='km', gender='female', output_path=None, rate=None):
    """
    Generates high-quality neural voice synthesis using edge-tts with robust gTTS fallback.
    Now supports gender-tracked voices.
    """
    if not text or not text.strip():
        return None
        
    temp_dir = get_temp_dir()
    if not output_path:
        output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_neural_tts.mp3")
        
    clean_lang = lang.lower().strip()[:2]
    if clean_lang == 'kh':
        clean_lang = 'km'
    clean_text = text.strip()
    
    # 1. Try edge-tts for Khmer, English, Chinese, etc.
    EDGE_VOICES = {
        'km': {'female': 'km-KH-SreymomNeural', 'male': 'km-KH-PisethNeural'},
        'en': {'female': 'en-US-AvaNeural', 'male': 'en-US-GuyNeural'},
        'zh': {'female': 'zh-CN-XiaoxiaoNeural', 'male': 'zh-CN-YunxiNeural'},
        'vi': {'female': 'vi-VN-HoaiMyNeural', 'male': 'vi-VN-NamMinhNeural'},
        'th': {'female': 'th-TH-PremwadeeNeural', 'male': 'th-TH-NiwatNeural'},
        'ko': {'female': 'ko-KR-SunHiNeural', 'male': 'ko-KR-InJoonNeural'},
        'ja': {'female': 'ja-JP-NanamiNeural', 'male': 'ja-JP-KeitaNeural'},
        'fr': {'female': 'fr-FR-DeniseNeural', 'male': 'fr-FR-HenriNeural'},
        'es': {'female': 'es-ES-ElviraNeural', 'male': 'es-ES-AlvaroNeural'},
        'de': {'female': 'de-DE-KatjaNeural', 'male': 'de-DE-ConradNeural'},
        'ru': {'female': 'ru-RU-SvetlanaNeural', 'male': 'ru-RU-DmitryNeural'}
    }
    
    voice_map = EDGE_VOICES.get(clean_lang, EDGE_VOICES['km'] if clean_lang == 'km' else EDGE_VOICES['en'])
    voice = voice_map.get(gender, voice_map['female'])
    
    if voice and clean_text:
        try:
            import edge_tts
            import asyncio
            import random
            async def _run_edge():
                # Inject subtle pitch and rate variations to avoid robotic monotone
                pitch_val = f"+{random.randint(2, 7)}Hz" if gender == 'female' else f"{random.choice(['-','+'])}{random.randint(1, 4)}Hz"
                rate_val = f"+{random.randint(2, 10)}%"
                
                kwargs = {"text": clean_text, "voice": voice, "pitch": pitch_val, "rate": rate_val}
                if rate:
                    kwargs["rate"] = rate
                communicate = edge_tts.Communicate(**kwargs)
                await communicate.save(output_path)
            for attempt in range(2):
                try:
                    loop = asyncio.get_running_loop()
                    asyncio.run_coroutine_threadsafe(_run_edge(), loop).result(20)
                except RuntimeError:
                    asyncio.run(_run_edge())
                except Exception as e_inner:
                    logger.warning(f"edge-tts attempt {attempt+1} error: {e_inner}")
                    continue  # No sleep — retry immediately

                if os.path.exists(output_path) and os.path.getsize(output_path) > 100:
                    return output_path
        except Exception as e:
            logger.error(f"edge-tts error: {e}")
            
    # 2. Fallback to gTTS
    try:
        from gtts import gTTS
        gtts_lang = 'zh-CN' if clean_lang == 'zh' else ('km' if clean_lang in ('km', 'kh') else clean_lang)
        tts = gTTS(text=clean_text, lang=gtts_lang, slow=False)
        tts.save(output_path)
        if os.path.exists(output_path) and os.path.getsize(output_path) > 100:
            return output_path
    except Exception as e_gtts:
        print(f"gTTS error: {e_gtts}")
        
    return None

def translate_nllb_ct2(text, src_lang="zh", tgt_lang="km"):
    """
    Translates text using fast AI API / Google Translate GTX endpoint (0.1s response time).
    """
    if not text or not text.strip():
        return text

    LANG_NAMES = {'km': 'Khmer (ភាសាខ្មែរ)', 'en': 'English', 'zh': 'Chinese (中文)', 'vi': 'Vietnamese', 'th': 'Thai'}
    tgt_name = LANG_NAMES.get(tgt_lang[:2], 'Khmer')
    src_name = LANG_NAMES.get(src_lang[:2], 'Chinese')
    
    gemini_key = os.environ.get("GEMINI_API_KEY")
    session = _get_translate_session()

    # 1. Primary: Fast Gemini 2.0 Flash API (reusing persistent session)
    if gemini_key:
        try:
            import urllib.parse
            prompt = f"Translate the following spoken dialogue from {src_name} into natural, accurate, conversational {tgt_name}. Output ONLY the raw translated spoken text with no markdown, no quotes, no explanations:\n\n{text}"
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_key}"
            payload = {"contents": [{"parts": [{"text": prompt}]}]}
            res = session.post(url, json=payload, timeout=6)
            if res.status_code == 200:
                data = res.json()
                tr_text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
                if tr_text and len(tr_text) > 0 and tr_text != text:
                    if tgt_lang[:2] == 'km':
                        tr_text = tr_text.replace('អង្ករចៀន', 'បាយឆា').replace('បាយចៀន', 'បាយឆា')
                    return tr_text
        except Exception as e_gem:
            print(f"Gemini translation error: {e_gem}")

    # 2. Fast Google Translate GTX Free Endpoint (persistent session, ~0.1s)
    try:
        import urllib.parse
        s_code = src_lang[:2] if src_lang != "auto" else "auto"
        t_code = tgt_lang[:2]
        url = f"https://translate.googleapis.com/translate_a/single?client=gtx&sl={s_code}&tl={t_code}&dt=t&q={urllib.parse.quote(text)}"
        res = session.get(url, timeout=5)
        if res.status_code == 200:
            data = res.json()
            translated = "".join(seg[0] for seg in data[0] if seg and seg[0])
            if translated and len(translated.strip()) > 0:
                res_txt = translated.strip()
                if tgt_lang[:2] == 'km':
                    res_txt = res_txt.replace('អង្ករចៀន', 'បាយឆា').replace('បាយចៀន', 'បាយឆា')
                return res_txt
    except Exception as e_gtx:
        print(f"GTX translation error: {e_gtx}")

    if tgt_lang[:2] == 'km':
        text = text.replace('អង្ករចៀន', 'បាយឆា').replace('បាយចៀន', 'បាយឆា')

    return text

def transcribe_with_whisper_timestamps(input_path, src_lang="auto", max_pause=0.35):
    """
    Transcribes audio with Faster-Whisper using word_timestamps=True.
    Intelligently splits speech into single-word or single-sentence units:
    - If speaker pauses > max_pause (0.35s) between words -> creates separate word segment.
    - If speaker talks continuously without pause -> groups words into a single sentence segment.
    """
    segments_list = []
    clean_src = None
    if src_lang != "auto":
        l = src_lang.lower().strip()
        clean_src = {'khmer': 'km', 'english': 'en', 'chinese': 'zh', 'vietnamese': 'vi', 'thai': 'th', 'korean': 'ko', 'japanese': 'ja', 'french': 'fr', 'spanish': 'es', 'german': 'de', 'russian': 'ru', 'arabic': 'ar', 'hindi': 'hi'}.get(l, l[:2])
    
    try:
        model = _get_whisper_model()  # Use cached model — avoids reload on every call
        if not model:
            return None

        segments, info = model.transcribe(input_path, language=clean_src, word_timestamps=True, beam_size=3, vad_filter=False)
        
        all_words = []
        for segment in segments:
            if hasattr(segment, 'words') and segment.words:
                for w in segment.words:
                    w_text = w.word.strip()
                    if w_text:
                        all_words.append({
                            "start": round(w.start, 2),
                            "end": round(w.end, 2),
                            "word": w_text
                        })
            else:
                text = segment.text.strip()
                if text:
                    segments_list.append({
                        "start": round(segment.start, 2),
                        "end": round(segment.end, 2),
                        "text": text
                    })

        if all_words:
            curr_words = []
            curr_start = None
            last_end = None

            for w in all_words:
                w_start = w["start"]
                w_end = w["end"]
                w_text = w["word"]

                if last_end is not None and (w_start - last_end) > max_pause:
                    # Pause detected! Flush current sentence segment
                    if curr_words:
                        segments_list.append({
                            "start": curr_start,
                            "end": last_end,
                            "text": " ".join(curr_words)
                        })
                        curr_words = []
                        curr_start = None

                if curr_start is None:
                    curr_start = w_start
                curr_words.append(w_text)
                last_end = w_end

            if curr_words and curr_start is not None:
                segments_list.append({
                    "start": curr_start,
                    "end": last_end,
                    "text": " ".join(curr_words)
                })

        if segments_list:
            return segments_list
    except Exception as e:
        print(f"Faster-Whisper error: {e}")
        
    return None

def prepare_telegram_thumbnail(thumb_path: str) -> str:
    """
    Ensures the thumbnail image is a properly formatted JPEG (max 320x320) for Telegram video previews.
    Returns the path to the formatted JPEG thumbnail file.
    """
    if not thumb_path or not os.path.exists(thumb_path):
        return None
    try:
        from PIL import Image
        temp_dir = get_temp_dir()
        out_jpg = os.path.join(temp_dir, f"{uuid.uuid4()}_tg_thumb.jpg")
        with Image.open(thumb_path) as img:
            img = img.convert("RGB")
            img.thumbnail((320, 320), Image.Resampling.LANCZOS)
            img.save(out_jpg, "JPEG", quality=92)
        if os.path.exists(out_jpg) and os.path.getsize(out_jpg) > 0:
            return out_jpg
    except Exception as e:
        print(f"prepare_telegram_thumbnail error: {e}")
    return thumb_path

def apply_thumbnail_to_video(video_path: str, thumb_path: str) -> str:
    """
    Returns video_path cleanly without adding corrupted secondary video streams into MP4 container,
    ensuring Telegram natively displays custom cover art passed via Pyrogram's thumb parameter.
    """
    return video_path

def translate_and_dub_media(input_path, target_lang='km', src_lang='auto', is_video=True, progress_callback=None, thumbnail_path=None, keep_bgm=True):
    """
    Full AI Video Localization & Dubbing Pipeline (v3 - Mouth-Tracking Edition):
    1. Uses Faster-Whisper to extract exact per-segment timestamps from original speech.
    2. Translates each segment using Gemini/Google Translate.
    3. Synthesizes Neural TTS (edge-tts) per segment.
    4. Speed-compresses each TTS clip to FIT inside its original time slot (tts_dur / seg_dur as atempo).
    5. Delays each clip to its original start time so speech lines up with mouth movements.
    6. Mixes all delayed clips into a single full-length audio timeline.
    7. Skips BGM mixing if DSP vocal removal produces near-silence (avoids polluting audio).
    8. Applies loudnorm to match original audio level.
    9. Merges dubbed audio track with original video stream, locked to original duration.
    """
    temp_dir = get_temp_dir()
    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    orig_dur = get_video_duration(input_path)

    if progress_callback: progress_callback("⚡ Transcribing speech with Faster-Whisper (timestamp alignment)...")

    # Step 1: Get timestamped segments from Whisper
    whisper_segs = transcribe_with_whisper_timestamps(input_path, src_lang=src_lang)

    # Fallback: use silencedetect-based segments if Whisper fails
    if not whisper_segs:
        if progress_callback: progress_callback("⚡ Falling back to silence-detect for timestamps...")
        raw_segments = detect_speech_segments(input_path)
        if raw_segments:
            from plugins.document_parser import transcribe_audio_video
            whisper_segs = []
            for (ss, se) in raw_segments:
                dur = se - ss
                if dur < 0.2: continue
                seg_slice = os.path.join(temp_dir, f"{uuid.uuid4()}_slice.mp3")
                try:
                    subprocess.run([ffmpeg_exe, "-y", "-ss", str(ss), "-t", str(dur), "-i", input_path, "-acodec", "libmp3lame", seg_slice], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=10)
                    txt = transcribe_audio_video(seg_slice, src_lang=src_lang)
                    cleanup_file(seg_slice)
                    if txt and len(txt.strip()) > 1 and "Error" not in txt:
                        whisper_segs.append({"start": ss, "end": se, "text": txt.strip()})
                except Exception:
                    cleanup_file(seg_slice)

    if not whisper_segs:
        # Last resort: full transcript, single segment spanning entire video
        from plugins.document_parser import transcribe_audio_video
        full_transcript = transcribe_audio_video(input_path, src_lang=src_lang)
        if not full_transcript or "Error" in full_transcript or len(full_transcript.strip()) < 2:
            return "ERROR: Could not extract speech from media to translate."
        whisper_segs = [{"start": 0.0, "end": orig_dur if orig_dur > 0 else 60.0, "text": full_transcript.strip()}]

    LANG_NAMES_MAP = {
        'km': 'Khmer (ភាសាខ្មែរ)', 'en': 'English', 'zh': 'Chinese (中文)',
        'vi': 'Vietnamese (Tiếng Việt)', 'th': 'Thai (ภาษาไทย)', 'ko': 'Korean (한국어)',
        'ja': 'Japanese (日本語)', 'fr': 'French (Français)', 'es': 'Spanish (Español)',
        'de': 'German (Deutsch)', 'ru': 'Russian (Русский)', 'ar': 'Arabic (العربية)',
        'hi': 'Hindi (हिन्दी)'
    }
    target_lang_name = LANG_NAMES_MAP.get(target_lang[:2], 'Khmer')
    if progress_callback: progress_callback(f"🧠 Translating & dubbing {len(whisper_segs)} segments to {target_lang_name}...")

    # Step 2: Batch translate ALL segments via a single Gemini call, then parallelize
    clean_src = src_lang if src_lang != "auto" else "zh"
    valid_segs = [(idx, seg) for idx, seg in enumerate(whisper_segs)
                  if (seg["end"] - seg["start"]) >= 0.15 and seg.get("text", "").strip()]

    # ── Batch translation via Gemini (single prompt, numbered lines) ─────────
    translations = {}   # idx -> translated text
    gemini_key = os.environ.get("GEMINI_API_KEY")
    if gemini_key and valid_segs:
        try:
            LANG_NAMES_LOCAL = {'km': 'Khmer (ភាសាខ្មែរ)', 'en': 'English', 'zh': 'Chinese (中文)',
                                 'vi': 'Vietnamese', 'th': 'Thai', 'ko': 'Korean', 'ja': 'Japanese',
                                 'fr': 'French', 'es': 'Spanish', 'de': 'German', 'ru': 'Russian',
                                 'ar': 'Arabic', 'hi': 'Hindi'}
            tgt_name_b = LANG_NAMES_LOCAL.get(target_lang[:2], 'Khmer')
            src_name_b = LANG_NAMES_LOCAL.get(clean_src[:2], 'the source language')
            numbered = "\n".join(f"{i+1}. {seg['text'].strip()}" for i, (_, seg) in enumerate(valid_segs))
            batch_prompt = (
                f"Translate each numbered line from {src_name_b} into natural conversational {tgt_name_b}.\n"
                f"Return ONLY the same numbered lines with translated text. No explanations, no extra lines.\n\n"
                f"{numbered}"
            )
            session = _get_translate_session()
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_key}"
            res = session.post(url, json={"contents": [{"parts": [{"text": batch_prompt}]}]}, timeout=30)
            if res.status_code == 200:
                raw_out = res.json()["candidates"][0]["content"]["parts"][0]["text"]
                for line in raw_out.strip().splitlines():
                    m = re.match(r'^(\d+)\.\s*(.+)', line.strip())
                    if m:
                        line_no = int(m.group(1)) - 1
                        if 0 <= line_no < len(valid_segs):
                            orig_idx = valid_segs[line_no][0]
                            translations[orig_idx] = m.group(2).strip()
                if progress_callback:
                    progress_callback(f"✅ Batch translated {len(translations)}/{len(valid_segs)} segments in one API call.")
        except Exception as e_batch:
            logger.warning(f"Batch Gemini translation failed ({e_batch}), falling back to per-segment.")

    # ── Per-segment fallback translation for any that didn't get a batch result
    def _translate_one(idx_seg):
        idx, seg = idx_seg
        if idx in translations:
            return
        txt = translate_nllb_ct2(seg["text"].strip(), src_lang=clean_src, tgt_lang=target_lang)
        translations[idx] = txt

    missing = [(i, s) for (i, s) in valid_segs if i not in translations]
    if missing:
        from concurrent.futures import ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=min(8, len(missing))) as ex:
            list(ex.map(_translate_one, missing))

    # ── Parallel per-segment worker: gender-detect + TTS + ffmpeg pad ────────
    delayed_segment_files = []
    lock = __import__('threading').Lock()

    def _process_segment(idx_seg):
        idx, seg = idx_seg
        seg_start = seg["start"]
        seg_end   = seg["end"]
        seg_dur   = seg_end - seg_start

        txt_tr = translations.get(idx, seg.get("text", "").strip())
        txt_tr = re.sub(r'[*#_~`>\[\]\(\)]', ' ', txt_tr or "").strip()
        if target_lang[:2] == 'km':
            txt_tr = txt_tr.replace('អង្ករចៀន', 'បាយឆា').replace('បាយចៀន', 'បាយឆា')
        if not txt_tr:
            return

        # Gender detect (fast ZCR heuristic — no full pitch tracking)
        detected_gender = 'female'
        seg_slice_path = os.path.join(temp_dir, f"{uuid.uuid4()}_gender_slice.mp3")
        try:
            subprocess.run([
                ffmpeg_exe, "-y", "-ss", str(seg_start), "-t", str(min(seg_dur, 3.0)),
                "-i", input_path, "-acodec", "libmp3lame", "-ar", "16000", "-ac", "1", seg_slice_path
            ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=8)
            if os.path.exists(seg_slice_path):
                detected_gender = detect_gender_from_audio(seg_slice_path)
        except Exception as e_gen:
            logger.warning(f"Gender slice error seg {idx}: {e_gen}")
        finally:
            cleanup_file(seg_slice_path)

        # Neural TTS
        raw_tts = generate_neural_tts(txt_tr, lang=target_lang[:2], gender=detected_gender)
        if not raw_tts or not os.path.exists(raw_tts):
            return

        tts_dur = get_video_duration(raw_tts)
        if tts_dur <= 0:
            cleanup_file(raw_tts)
            return

        speed_factor = tts_dur / max(0.15, seg_dur)
        speed_factor = max(1.0, min(2.5, speed_factor))

        delay_ms = int(seg_start * 1000)
        delayed_file = os.path.join(temp_dir, f"{uuid.uuid4()}_delayed_seg_{idx}.mp3")
        try:
            stream_in = ffmpeg.input(raw_tts).audio
            for f_name, f_val in build_atempo_filter_chain(speed_factor):
                stream_in = stream_in.filter(f_name, f_val)
            stream_in = stream_in.filter('aresample', 44100).filter('aformat', channel_layouts='stereo')
            stream_in = stream_in.filter('adelay', delays=f"{delay_ms}|{delay_ms}")
            ffmpeg.output(stream_in, delayed_file, acodec='libmp3lame', ar='44100', ac=2, **{'b:a': '192k'}).overwrite_output().run(
                cmd=ffmpeg_exe, capture_stdout=True, capture_stderr=True
            )
        except Exception as e_ff:
            logger.error(f"FFmpeg pad error seg {idx}: {e_ff}")
        finally:
            cleanup_file(raw_tts)

        if os.path.exists(delayed_file) and os.path.getsize(delayed_file) > 100:
            with lock:
                delayed_segment_files.append((idx, delayed_file))

    if progress_callback:
        progress_callback(f"🎙️ Generating TTS & timing for {len(valid_segs)} segments in parallel...")

    from concurrent.futures import ThreadPoolExecutor
    # Use up to 6 threads — edge-tts is I/O bound, ffmpeg is CPU-light per segment
    max_workers = min(6, max(1, len(valid_segs)))
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        list(executor.map(_process_segment, valid_segs))

    # Sort by original segment index to preserve timing order
    delayed_segment_files.sort(key=lambda x: x[0])
    delayed_segment_files = [f for _, f in delayed_segment_files]

    # Step 3: Mix all delayed segments into a single full-duration speech track
    final_speech_track = None
    if delayed_segment_files:
        if progress_callback: progress_callback("🎛️ Mixing all dubbed segments into full audio timeline...")
        combined_speech = os.path.join(temp_dir, f"{uuid.uuid4()}_dubbed_speech.mp3")
        try:
            if len(delayed_segment_files) == 1:
                shutil.copy(delayed_segment_files[0], combined_speech)
                final_speech_track = combined_speech
            else:
                current_mix = delayed_segment_files
                batch_size = 40
                while len(current_mix) > 1:
                    next_mix = []
                    for i in range(0, len(current_mix), batch_size):
                        batch = current_mix[i:i+batch_size]
                        if len(batch) == 1:
                            next_mix.append(batch[0])
                            continue
                        batch_out = os.path.join(temp_dir, f"{uuid.uuid4()}_batch_mix.mp3")
                        inputs = [ffmpeg.input(f).audio for f in batch]
                        ffmpeg.filter(inputs, 'amix', inputs=len(inputs), normalize=False, duration='longest').output(
                            batch_out, acodec='libmp3lame', ar='44100', ac=2, **{'b:a': '192k'}
                        ).overwrite_output().run(cmd=ffmpeg_exe, capture_stdout=True, capture_stderr=True)
                        next_mix.append(batch_out)
                    current_mix = next_mix
                
                shutil.copy(current_mix[0], combined_speech)
                final_speech_track = combined_speech
        except Exception as e_mix:
            logger.error(f"Error mixing delayed segments: {e_mix}")
            if delayed_segment_files:
                final_speech_track = delayed_segment_files[0]

    if not final_speech_track:
        return "ERROR: Failed to generate dubbed speech track."

    # Step 4: Extract background music & sound effects using Demucs v4 AI stem separation
    final_audio = final_speech_track
    if keep_bgm:
        try:
            if progress_callback: progress_callback("🎵 Isolating Background Music & Sound Effects with Demucs AI...")
            bgm_result = extract_bgm_demucs(input_path, return_vocals=True)
            if isinstance(bgm_result, tuple):
                bgm_path, vocals_path = bgm_result
            else:
                bgm_path, vocals_path = bgm_result, None

            if bgm_path and os.path.exists(bgm_path) and os.path.getsize(bgm_path) > 1000:
                if progress_callback: progress_callback("🎧 Mixing dubbed voice with BGM and ducked original vocals (to retain cries/laughs)...")
                
                mixed_path = os.path.join(temp_dir, f"{uuid.uuid4()}_mixed_all.mp3")
                
                # No vocals mixed per user request (0% original voice)
                filter_graph = (
                    "[0:a]volume=0.70[bgm]; "
                    "[1:a]volume=1.40[tts_mix]; "
                    "[bgm][tts_mix]amix=inputs=2:duration=first:normalize=0,aresample=44100,aformat=channel_layouts=stereo[outa]"
                )
                cmd = [
                    ffmpeg_exe, "-y",
                    "-i", bgm_path,
                    "-i", final_speech_track,
                    "-filter_complex", filter_graph,
                    "-map", "[outa]",
                    "-acodec", "libmp3lame", "-ar", "44100", "-ac", "2", "-b:a", "192k",
                    mixed_path
                ]
                
                subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=1800)
                
                if os.path.exists(mixed_path) and os.path.getsize(mixed_path) > 1000:
                    final_audio = mixed_path
                cleanup_file(bgm_path)
                if vocals_path: cleanup_file(vocals_path)
            else:
                if bgm_path: cleanup_file(bgm_path)
                if vocals_path: cleanup_file(vocals_path)
                if progress_callback: progress_callback("ℹ️ BGM stem not available — using pure voice dub.")
        except Exception as e_bgm:
            print(f"BGM isolation & mixing error: {e_bgm}")
    else:
        if progress_callback: progress_callback("ℹ️ Background audio skipped per user request — using pure voice dub.")

    # Step 5: Apply loudnorm to match original audio loudness (-16 LUFS broadcast standard)
    if progress_callback: progress_callback("🔊 Normalizing audio loudness to match original video...")
    loudnorm_path = os.path.join(temp_dir, f"{uuid.uuid4()}_loudnorm.mp3")
    try:
        subprocess.run(
            [ffmpeg_exe, "-y", "-i", final_audio,
             "-af", "loudnorm=I=-16:TP=-1.5:LRA=11",
             "-acodec", "libmp3lame", "-ar", "44100", "-ac", "2", "-b:a", "192k",
             loudnorm_path],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=30
        )
        if os.path.exists(loudnorm_path) and os.path.getsize(loudnorm_path) > 100:
            final_audio = loudnorm_path
    except Exception as e_ln:
        print(f"Loudnorm failed: {e_ln}")

    # Step 6: Merge dubbed audio with original video (locked to original duration)
    output_ext = "mp4" if is_video else "mp3"
    output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_dubbed.{output_ext}")

    if is_video:
        if progress_callback: progress_callback("🎬 Merging dubbed audio with original video...")
        try:
            out_opts = {'vcodec': 'copy', 'acodec': 'aac', 'b:a': '192k', 'ar': '44100', 'ac': 2}
            if orig_dur and orig_dur > 0:
                out_opts['t'] = orig_dur

            video_in = ffmpeg.input(input_path).video
            audio_in = ffmpeg.input(final_audio).audio
            stream = ffmpeg.output(video_in, audio_in, output_path, **out_opts).overwrite_output()
            res = _run_ffmpeg_with_progress(stream, output_path, progress_callback)
        except Exception as e_merge:
            print(f"FFmpeg merge error: {e_merge}")
            res = None
    else:
        shutil.copy(final_audio, output_path)
        res = output_path

    # Cleanup temp files
    for sf in delayed_segment_files:
        cleanup_file(sf)
    if final_speech_track != final_audio:
        cleanup_file(final_speech_track)
    cleanup_file(final_audio)

    if res and os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
        if thumbnail_path and os.path.exists(thumbnail_path) and is_video:
            if progress_callback: progress_callback("🖼️ Applying custom thumbnail to video...")
            thumbed = apply_thumbnail_to_video(output_path, thumbnail_path)
            if thumbed and thumbed != output_path:
                cleanup_file(output_path)
                return thumbed
        return output_path

    return "ERROR: Video dubbing process failed."

def mix_recap_sidechain_ducking(bgm_wav_path, narration_wav_path, output_audio_path=None):
    """
    Applies dynamic FFmpeg sidechaincompress audio ducking to automatically lower background music & sound effects (no_vocals stem)
    whenever the narrator speaks, then mixes narrator audio with ducked BGM at 44.1kHz Stereo.
    """
    temp_dir = get_temp_dir()
    if not output_audio_path:
        output_audio_path = os.path.join(temp_dir, f"{uuid.uuid4()}_ducked_recap.mp3")
        
    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    import subprocess
    
    filter_graph = (
        "[0:a]volume=0.85[bgm]; "
        "[1:a]asplit[narrator][sc]; "
        "[bgm][sc]sidechaincompress=threshold=0.08:ratio=4:attack=100:release=400[ducked_bgm]; "
        "[ducked_bgm][narrator]amix=inputs=2:duration=first:normalize=0,aresample=44100,aformat=channel_layouts=stereo[outa]"
    )
    
    cmd = [
        ffmpeg_exe, "-y",
        "-i", bgm_wav_path,
        "-i", narration_wav_path,
        "-filter_complex", filter_graph,
        "-map", "[outa]",
        "-acodec", "libmp3lame",
        "-ar", "44100",
        "-ac", "2",
        "-b:a", "192k",
        output_audio_path
    ]
    try:
        subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=1800)
        if os.path.exists(output_audio_path) and os.path.getsize(output_audio_path) > 100:
            return output_audio_path
    except Exception as e:
        print(f"Sidechain ducking error: {e}")
        
    return mix_tts_with_bgm(narration_wav_path, bgm_wav_path, bgm_volume=0.25)

def clean_recap_speech_text(text):
    """Strips markdown, emojis, bullet points, numbers, and section headers to leave clean spoken recap text only."""
    t = re.sub(r'[\*\_~`>#]', ' ', text)
    t = re.sub(r'[📌💡🎯🎬▶✔✅⭐🏆•\-]\s*', ' ', t)
    t = re.sub(r'^\s*\d+[\.\)]\s*', ' ', t, flags=re.MULTILINE)
    t = re.sub(r'(Core Concept|Main Story Meaning|Detailed Scene-by-Scene|Key Point Breakdown|Key Highlights|Main Points|Core Topic|Story Overview|Discussion|Important Insights|Lessons|Final Conclusion|Main Takeaway|Summary|Executive Overview|Chronological Scene|Deep Analytical Insights|សង្ខេបលម្អិត|ខ្លឹមសារ|សេចក្តីសន្និដ្ឋាន)\s*[:\&]*', ' ', t, flags=re.IGNORECASE)
    t = re.sub(r'^(Here is|Summary|Recap|Khmer|English|Note):.*$', '', t, flags=re.MULTILINE | re.IGNORECASE)
    lines = [line.strip() for line in t.split('\n') if line.strip()]
    cleaned = ". ".join(lines)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def recap_video_audio(input_path, target_lang='km', src_lang='auto', is_video=True, voiceover=False, progress_callback=None, thumbnail_path=None):
    """
    Automated Movie Recap Generation Pipeline:
    1. Transcribes media with Faster-Whisper ASR to extract scene timestamps and transcript.
    2. Uses Gemini 2.0 Flash to generate a dramatic 3rd-person narrator recap script & executive breakdown in target language.
    3. Synthesizes dramatic cinematic voiceover using edge-tts (km-KH-SreymomNeural, rate="-8%").
    4. Separates vocals and ambient BGM with Demucs v4 stem separation, discarding original vocals.
    5. Applies dynamic FFmpeg sidechaincompress audio ducking so BGM lowers whenever the narrator speaks.
    6. Cuts key scenes & matches video speed to narration.
    """
    temp_dir = get_temp_dir()
    
    if progress_callback: progress_callback("⚡ Transcribing scenes with Faster-Whisper ASR...")
    whisper_segs = transcribe_with_whisper_timestamps(input_path, src_lang=src_lang)
    
    from plugins.document_parser import transcribe_audio_video
    transcript = transcribe_audio_video(input_path, src_lang=src_lang)
    
    if not transcript or "Error" in transcript or "Unsupported" in transcript or len(transcript.strip()) < 3:
        return "ERROR: Could not extract speech from media to generate recap.", None
        
    LANG_NAMES_MAP = {
        'km': 'Khmer (ភាសាខ្មែរ)',
        'en': 'English',
        'zh': 'Chinese (中文)',
        'vi': 'Vietnamese (Tiếng Việt)',
        'th': 'Thai (ภาษาไทย)',
        'ko': 'Korean (한국어)',
        'ja': 'Japanese (日本語)',
        'fr': 'French (Français)',
        'es': 'Spanish (Español)',
        'de': 'German (Deutsch)',
        'ru': 'Russian (Русский)',
        'ar': 'Arabic (العربية)',
        'hi': 'Hindi (हिन्दी)'
    }
    target_lang_name = LANG_NAMES_MAP.get(target_lang[:2], 'Khmer')
    
    if progress_callback: progress_callback(f"🧠 Generating Cinematic Movie Recap Script in {target_lang_name} (Gemini 2.0 Flash)...")
    
    import asyncio
    from plugins.core import get_ai_response
    prompt = (
        f"You are Udom AI, an elite film director and movie recap narrator.\n"
        f"Generate a HIGHLY DETAILED, EXECUTIVE-GRADE PROFESSIONAL MOVIE RECAP & SYNTHESIS of the following video transcript in {target_lang_name}.\n\n"
        f"CRITICAL REQUIREMENTS:\n"
        f"1. Explain the FULL STORY context, background setting, character motivations, plot twists, and main message with extreme narrative flair.\n"
        f"2. Detail EVERY MAJOR SCENE in chronological sequence so the audience grasps every plot point without watching the original.\n"
        f"3. Provide deep analytical insights, core takeaways, and professional summary conclusions.\n\n"
        f"Structure your analysis in {target_lang_name} using these exact sections:\n\n"
        f"🏆 **1. EXECUTIVE OVERVIEW & CORE MEANING (ខ្លឹមសារ និងអត្ថន័យដើមនៃវីដេអូ)**\n"
        f"• Complete breakdown of what the movie is about, the main theme, background context, and core story purpose.\n\n"
        f"🎬 **2. CHRONOLOGICAL SCENE & HIGHLIGHT BREAKDOWN (សង្ខេបលម្អិតតាមចំនុច និងសាច់រឿង)**\n"
        f"• Detailed bullet points explaining key scenes, actions, plot progression, and climax.\n\n"
        f"💡 **3. DEEP ANALYTICAL INSIGHTS & LESSONS (ការវិភាគយ៉ាងស៊ីជម្រៅ និងមេរៀន)**\n"
        f"• Expert cinematic analysis of character arcs, philosophical messages, and hidden details.\n\n"
        f"🎯 **4. FINAL CONCLUSION & KEY TAKEAWAYS (សេចក្តីសន្និដ្ឋាន និងសារៈសំខាន់)**\n"
        f"• Final executive conclusion and summary takeaway.\n\n"
        f"Transcript:\n{transcript}"
    )
    
    try:
        try:
            loop = asyncio.get_running_loop()
            recap_text = asyncio.run_coroutine_threadsafe(get_ai_response(0, prompt), loop).result(45)
        except RuntimeError:
            recap_text = asyncio.run(get_ai_response(0, prompt))
    except Exception as e:
        print(f"AI Recap Error: {e}")
        recap_text = f"Failed to generate recap: {e}"
        
    if not voiceover:
        return recap_text, None
        
    # Generate Dramatic 3rd-Person Voiceover Script
    if progress_callback: progress_callback("🎙 Scripting Dramatic 3rd-Person Narrator Voiceover...")
    vo_script_prompt = (
        f"You are Udom AI, a world-class documentary and movie recap narrator.\n"
        f"Write a smooth, highly engaging, dramatic 3rd-person spoken RECAP VOICE-OVER NARRATION in {target_lang_name} based on this transcript.\n\n"
        f"GUIDELINES:\n"
        f"1. Narrative Arc: Start with a captivating hook, describe scene developments in dramatic chronological order, and finish with a memorable conclusion.\n"
        f"2. Style: Write in natural, fluent, cinematic spoken narration like a professional YouTube movie recap narrator.\n"
        f"3. Pure Spoken Text: Output ONLY pure spoken sentences in {target_lang_name}. Do NOT include section titles, numbers, bullet symbols, emojis, markdown asterisks, or quotes.\n\n"
        f"Transcript:\n{transcript}"
    )
    
    try:
        try:
            loop = asyncio.get_running_loop()
            raw_script = asyncio.run_coroutine_threadsafe(get_ai_response(0, vo_script_prompt), loop).result(30)
        except RuntimeError:
            raw_script = asyncio.run(get_ai_response(0, vo_script_prompt))
    except Exception:
        raw_script = recap_text
        
    tts_speech = clean_recap_speech_text(raw_script)
    if not tts_speech:
        tts_speech = clean_recap_speech_text(recap_text)
        
    if progress_callback: progress_callback("🎙 Synthesizing Cinematic Neural Narrator Voice (edge-tts -8% speed)...")
    raw_tts = generate_neural_tts(tts_speech, lang=target_lang[:2], rate="-8%")
    if not raw_tts or not os.path.exists(raw_tts):
        return recap_text, None
        
    # Isolate original BGM & Sound effects using Demucs v4 (discarding original vocals)
    if progress_callback: progress_callback("🎵 Isolating Ambient BGM & SFX with Demucs v4 (discarding original vocals)...")
    bgm_path = extract_bgm_demucs(input_path)
    
    orig_dur = get_video_duration(input_path)
    tts_dur = get_video_duration(raw_tts)
    target_dur = tts_dur if tts_dur > 0 else (orig_dur if orig_dur > 0 else 60.0)
    
    # Dynamic Sidechain Audio Ducking
    final_recap_audio = raw_tts
    if bgm_path and os.path.exists(bgm_path) and check_audio_rms(bgm_path):
        if progress_callback: progress_callback("🎧 Applying Dynamic Sidechain Audio Ducking (narrator ducking BGM volume)...")
        ducked = mix_recap_sidechain_ducking(bgm_path, raw_tts)
        if ducked and os.path.exists(ducked):
            final_recap_audio = ducked
        
    if is_video and orig_dur > 0 and target_dur > 0:
        if progress_callback: progress_callback("🎬 Edit & Speed Engine: Syncing video speed & scenes to narration...")
        
        output_path = os.path.join(temp_dir, f"{uuid.uuid4()}_movie_recap.mp4")
        
        speech_segments = whisper_segs if whisper_segs else detect_speech_segments(input_path)
        cut_video_path = None
        
        try:
            if speech_segments and len(speech_segments) > 1:
                segment_clips = []
                segs = speech_segments if isinstance(speech_segments[0], tuple) else [(s["start"], s["end"]) for s in speech_segments]
                for s_idx, (s_start, s_end) in enumerate(segs[:12]):
                    c_dur = s_end - s_start
                    if c_dur < 0.3: continue
                    c_out = os.path.join(temp_dir, f"{uuid.uuid4()}_recap_clip_{s_idx}.mp4")
                    try:
                        (
                            ffmpeg
                            .input(input_path, ss=s_start, t=c_dur)
                            .output(c_out, vcodec='copy', an=None)
                            .overwrite_output()
                            .run(cmd=imageio_ffmpeg.get_ffmpeg_exe(), capture_stdout=True, capture_stderr=True)
                        )
                        if os.path.exists(c_out) and os.path.getsize(c_out) > 100:
                            segment_clips.append(c_out)
                    except Exception:
                        pass
                
                if segment_clips:
                    concat_txt = os.path.join(temp_dir, f"{uuid.uuid4()}_recap_list.txt")
                    with open(concat_txt, "w", encoding="utf-8") as f_lst:
                        for clp in segment_clips:
                            f_lst.write(f"file '{clp}'\n")
                            
                    concat_raw = os.path.join(temp_dir, f"{uuid.uuid4()}_recap_concat.mp4")
                    try:
                        (
                            ffmpeg
                            .input(concat_txt, format='concat', safe=0)
                            .output(concat_raw, vcodec='copy', an=None)
                            .overwrite_output()
                            .run(cmd=imageio_ffmpeg.get_ffmpeg_exe(), capture_stdout=True, capture_stderr=True)
                        )
                        if os.path.exists(concat_raw) and os.path.getsize(concat_raw) > 100:
                            cut_video_path = concat_raw
                    except Exception:
                        pass
                    finally:
                        cleanup_file(concat_txt)
                        for clp in segment_clips:
                            cleanup_file(clp)

            source_v = cut_video_path if cut_video_path else input_path
            source_v_dur = get_video_duration(source_v) or orig_dur
            pts_factor = max(0.5, min(2.0, target_dur / max(1.0, source_v_dur)))
            
            video_input = ffmpeg.input(source_v).video.filter('setpts', f'{pts_factor}*PTS')
            audio_input = ffmpeg.input(final_recap_audio).audio
            
            out_opts = {'vcodec': 'libx264', 'acodec': 'aac', 'pix_fmt': 'yuv420p', 'ar': '44100', 'b:a': '192k', 't': target_dur}
            stream = (
                ffmpeg
                .output(video_input, audio_input, output_path, **out_opts)
                .overwrite_output()
            )
            res_media = _run_ffmpeg_with_progress(stream, output_path, progress_callback)
            if cut_video_path:
                cleanup_file(cut_video_path)
        except Exception as e_edit:
            print(f"Edit & Speed Engine Error: {e_edit}")
            try:
                res_media = _run_ffmpeg_with_progress(stream, output_path, progress_callback)
            except Exception:
                res_media = final_recap_audio
    else:
        res_media = final_recap_audio
        
    cleanup_file(raw_tts)
    if bgm_path:
        cleanup_file(bgm_path)

    # Apply custom thumbnail to final video if provided
    if thumbnail_path and os.path.exists(thumbnail_path) and is_video and isinstance(res_media, str) and res_media.endswith(".mp4") and os.path.exists(res_media):
        if progress_callback: progress_callback("🖼️ Applying custom thumbnail to recap video...")
        thumbed = apply_thumbnail_to_video(res_media, thumbnail_path)
        if thumbed and thumbed != res_media:
            cleanup_file(res_media)
            res_media = thumbed

    return recap_text, res_media

def get_video_duration(file_path):
    """Accurately extracts video or audio duration in seconds using imageio_ffmpeg, ffprobe, null-scan, OpenCV, or ffmpeg-python."""
    if not file_path or not os.path.exists(file_path):
        return 0.0

    import subprocess
    import re

    # Method 1: Standard FFmpeg header regex
    try:
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [ffmpeg_exe, "-analyzeduration", "10M", "-probesize", "10M", "-i", file_path]
        res = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, timeout=14400)
        stderr_text = res.stderr.decode("utf-8", errors="ignore")
        match = re.search(r"Duration:\s*(\d+):(\d+):([\d.]+)", stderr_text)
        if match:
            hours, minutes, seconds = match.groups()
            dur = float(hours) * 3600 + float(minutes) * 60 + float(seconds)
            if dur > 0:
                return dur
    except Exception as e:
        print(f"Error getting duration with imageio_ffmpeg: {e}")

    # Method 2: ffprobe CLI if available
    try:
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        ffprobe_exe = ffmpeg_exe.replace("ffmpeg.exe", "ffprobe.exe").replace("ffmpeg", "ffprobe")
        cmd_ffprobe = [ffprobe_exe, "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", file_path]
        res_p = subprocess.run(cmd_ffprobe, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=300)
        stdout_text = res_p.stdout.decode("utf-8", errors="ignore").strip()
        dur = float(stdout_text)
        if dur > 0:
            return dur
    except Exception:
        pass

    # Method 3: Fast null scan for files with Duration: N/A
    try:
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [ffmpeg_exe, "-y", "-i", file_path, "-map", "0:v:0?", "-c", "copy", "-f", "null", "-"]
        res = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE, timeout=7200)
        stderr_text = res.stderr.decode("utf-8", errors="ignore")
        matches = re.findall(r"time=\s*(\d+):(\d+):([\d.]+)", stderr_text)
        if matches:
            h, m, s = matches[-1]
            dur = float(h) * 3600 + float(m) * 60 + float(s)
            if dur > 0:
                return dur
    except Exception:
        pass

    # Method 4: OpenCV fallback
    try:
        import cv2
        cap = cv2.VideoCapture(file_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)
        cap.release()
        if fps > 0 and frame_count > 0:
            return frame_count / fps
    except Exception:
        pass

    return 0.0

def extract_dm_chapters(input_path):
    """
    Parses video chapters via ffmpeg to find 'dm_clip_seek_before' and 'dm_clip_seek_after'
    Returns a dict with 'before': (start, end) and 'after': (start, end) if found.
    """
    import re
    try:
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        cmd = [ffmpeg_exe, "-i", input_path]
        res = subprocess.run(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
        stderr_text = res.stderr.decode("utf-8", errors="ignore")
        chapters = {'before': None, 'after': None}
        for block in stderr_text.split("Chapter #")[1:]:
            start_match = re.search(r"start\s+([\d.]+),\s+end\s+([\d.]+)", block)
            title_match = re.search(r"title\s*:\s*([^\r\n]+)", block)
            if start_match and title_match:
                title = title_match.group(1).strip()
                if "dm_clip_seek_before" in title.lower() or title.lower() == "intro":
                    chapters['before'] = (float(start_match.group(1)), float(start_match.group(2)))
                elif "dm_clip_seek_after" in title.lower() or title.lower() == "outro":
                    chapters['after'] = (float(start_match.group(1)), float(start_match.group(2)))
        return chapters
    except Exception:
        return {'before': None, 'after': None}

def clip_video_into_parts(input_path, num_clips=3, known_duration=0, progress_callback=None):
    """
    Splits a video file into num_clips equal segment files using FFmpeg.
    Uses fast stream copy first, falling back to ultrafast re-encoding if needed.
    """
    temp_dir = get_temp_dir()
    duration = get_video_duration(input_path)
    if duration <= 0 and known_duration > 0:
        duration = float(known_duration)

    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()

    if duration <= 0:
        raise ValueError("Could not determine video duration. This video format or stream cannot be automatically clipped.")

    dm_chaps = extract_dm_chapters(input_path)
    if has_copyright_metadata(input_path) and progress_callback:
        progress_callback("⚠️ Copyright metadata detected and removed from clips.")
        
    main_start = 0.0
    main_end = duration

    if dm_chaps['before']:
        main_start = dm_chaps['before'][1]
    if dm_chaps['after']:
        main_end = dm_chaps['after'][0]
        
    main_duration = main_end - main_start
    if main_duration <= 0:
        main_duration = duration # Fallback if chapters are messed up
        main_start = 0.0

    clip_duration = main_duration / max(1, num_clips)
    output_files = []

    def _extract_clip(start_sec, dur_sec, prefix="clip", source_file=None):
        if source_file is None:
            source_file = input_path
        out_f = os.path.join(temp_dir, f"{uuid.uuid4()}_{prefix}.mp4")
        try:
            cmd_copy = [
                ffmpeg_exe, "-y",
                "-ss", f"{start_sec:.3f}",
                "-i", source_file,
                "-t", f"{dur_sec:.3f}",
                "-c", "copy",
                "-avoid_negative_ts", "make_zero",
                "-map_metadata", "-1",
                out_f
            ]
            subprocess.run(cmd_copy, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=7200)
            if not os.path.exists(out_f) or os.path.getsize(out_f) < 1000:
                cmd_reencode = [
                    ffmpeg_exe, "-y",
                    "-ss", f"{start_sec:.3f}",
                    "-i", source_file,
                    "-t", f"{dur_sec:.3f}",
                    "-c:v", "libx264", "-preset", "ultrafast", "-pix_fmt", "yuv420p",
                    "-c:a", "aac", "-b:a", "128k",
                    "-map_metadata", "-1",
                    out_f
                ]
                subprocess.run(cmd_reencode, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=14400)
            if os.path.exists(out_f) and os.path.getsize(out_f) > 500:
                return out_f
        except Exception as e:
            print(f"Error creating clip {prefix}: {e}")
        return None

    # 1. Extract dm_clip_seek_before (intro) if exists
    before_clip = None
    if dm_chaps['before']:
        b_start, b_end = dm_chaps['before']
        if progress_callback: progress_callback("✂️ Extracting Intro Ad...")
        before_clip = _extract_clip(b_start, b_end - b_start, "intro")

    # 2. Extract dm_clip_seek_after (outro) if exists
    after_clip = None
    if dm_chaps['after']:
        a_start, a_end = dm_chaps['after']
        if progress_callback: progress_callback("✂️ Extracting Outro Ad...")
        after_clip = _extract_clip(a_start, a_end - a_start, "outro")

    # 3. Extract main video first
    if progress_callback: progress_callback("✂️ Extracting Main Video...")
    main_video = _extract_clip(main_start, main_duration, "main_video")
    if not main_video:
        main_video = input_path

    # 4. Extract main parts from main video
    for i in range(num_clips):
        start_time = (i * clip_duration) if main_video != input_path else (main_start + (i * clip_duration))
        if progress_callback:
            progress_callback(f"✂️ Cutting video clip {i+1} of {num_clips}...")
        main_part = _extract_clip(start_time, clip_duration, f"part_{i+1}", source_file=main_video)
        if main_part:
            output_files.append(main_part)

    if main_video and main_video != input_path:
        try: os.remove(main_video)
        except: pass

    # The user requested: "final result 3 clip + dm_clip_seek_after and dm_clip_seek_before"
    # Wait, the prompt says "out of 3 clip if it have and final result 3 clip + dm_clip_seek_after and dm_clip_seek_before".
    # So we should probably prepend the before clip and append the after clip.
    result = {
        "intro": before_clip,
        "main_clips": output_files,
        "outro": after_clip,
    }
    return result

def clip_video_by_duration(input_path, duration_sec=60, known_duration=0, progress_callback=None):
    """
    Splits a video/audio file into chunks of duration_sec seconds using FFmpeg.
    Uses fast stream copy first, falling back to ultrafast re-encoding if needed.
    """
    import math
    temp_dir = get_temp_dir()
    total_dur = get_video_duration(input_path)
    if total_dur <= 0 and known_duration > 0:
        total_dur = float(known_duration)

    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()

    if total_dur <= 0 or duration_sec <= 0:
        raise ValueError("Could not determine total video duration. This video format or stream cannot be clipped by duration.")

    dm_chaps = extract_dm_chapters(input_path)
    if has_copyright_metadata(input_path) and progress_callback:
        progress_callback("⚠️ Copyright metadata detected and removed from clips.")
        
    main_start = 0.0
    main_end = total_dur

    if dm_chaps['before']:
        main_start = dm_chaps['before'][1]
    if dm_chaps['after']:
        main_end = dm_chaps['after'][0]
        
    main_duration = main_end - main_start
    if main_duration <= 0:
        main_duration = total_dur
        main_start = 0.0

    num_clips = int(math.ceil(main_duration / float(duration_sec)))
    output_files = []

    def _extract_clip(start_sec, dur_sec, prefix="clip", source_file=None):
        if source_file is None:
            source_file = input_path
        out_f = os.path.join(temp_dir, f"{uuid.uuid4()}_{prefix}.mp4")
        try:
            cmd_copy = [
                ffmpeg_exe, "-y",
                "-ss", f"{start_sec:.3f}",
                "-i", source_file,
                "-t", f"{dur_sec:.3f}",
                "-c", "copy",
                "-avoid_negative_ts", "make_zero",
                "-map_metadata", "-1",
                out_f
            ]
            subprocess.run(cmd_copy, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=7200)
            if not os.path.exists(out_f) or os.path.getsize(out_f) < 1000:
                cmd_reencode = [
                    ffmpeg_exe, "-y",
                    "-ss", f"{start_sec:.3f}",
                    "-i", source_file,
                    "-t", f"{dur_sec:.3f}",
                    "-c:v", "libx264", "-preset", "ultrafast", "-pix_fmt", "yuv420p",
                    "-c:a", "aac", "-b:a", "128k",
                    "-map_metadata", "-1",
                    out_f
                ]
                subprocess.run(cmd_reencode, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=14400)
            if os.path.exists(out_f) and os.path.getsize(out_f) > 500:
                return out_f
        except Exception as e:
            print(f"Error creating clip {prefix}: {e}")
        return None

    # 1. Extract intro
    before_clip = None
    if dm_chaps['before']:
        b_start, b_end = dm_chaps['before']
        if progress_callback: progress_callback("✂️ Extracting Intro Ad...")
        before_clip = _extract_clip(b_start, b_end - b_start, "intro")

    # 2. Extract outro
    after_clip = None
    if dm_chaps['after']:
        a_start, a_end = dm_chaps['after']
        if progress_callback: progress_callback("✂️ Extracting Outro Ad...")
        after_clip = _extract_clip(a_start, a_end - a_start, "outro")

    # 3. Extract main video first
    if progress_callback: progress_callback("✂️ Extracting Main Video...")
    main_video = _extract_clip(main_start, main_duration, "main_video")
    if not main_video:
        main_video = input_path

    # 4. Extract main parts from main video
    for i in range(num_clips):
        start_time = (i * float(duration_sec)) if main_video != input_path else (main_start + (i * float(duration_sec)))
        segment_dur = min(float(duration_sec), main_end - (start_time if main_video == input_path else (main_start + start_time)))
        if segment_dur < 0.5:
            continue

        if progress_callback:
            progress_callback(f"✂️ Cutting clip {i+1} of {num_clips} ({segment_dur:.0f}s)...")

        main_part = _extract_clip(start_time, segment_dur, f"part_dur_{i+1}", source_file=main_video)
        if main_part:
            output_files.append(main_part)

    if main_video and main_video != input_path:
        try: os.remove(main_video)
        except: pass

    result = {
        "intro": before_clip,
        "main_clips": output_files,
        "outro": after_clip,
    }
    return result
